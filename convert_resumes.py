#!/usr/bin/env python3
"""
Resume Converter - Converts Markdown resumes to ATS-friendly DOCX format.

DOCX is the preferred format for most Applicant Tracking Systems (ATS).
PDF can sometimes lose formatting or fail to parse in ATS systems.
DOCX preserves text structure and is consistently machine-readable.

Usage:
    # Convert a single resume:
    python3 convert_resumes.py Resumes/Jane_Doe_Resume_Acme_Corp.md

    # Convert all resumes:
    python3 convert_resumes.py --all

    # Convert to a specific output directory:
    python3 convert_resumes.py --all --output Output/
"""

import sys
import os
import re
import argparse
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import markdown2
from weasyprint import HTML


def parse_markdown(md_text):
    """Parse markdown into structured sections."""
    lines = md_text.strip().split('\n')
    sections = []
    current_section = None
    current_content = []

    for line in lines:
        if line.startswith('# '):
            if current_section:
                sections.append((current_section, current_content))
            current_section = ('h1', line[2:].strip())
            current_content = []
        elif line.startswith('## '):
            if current_section:
                sections.append((current_section, current_content))
            current_section = ('h2', line[3:].strip())
            current_content = []
        elif line.startswith('### '):
            current_content.append(('h3', line[4:].strip()))
        elif line.startswith('- **') or line.startswith('- '):
            current_content.append(('bullet', line[2:].strip()))
        elif line.startswith('---'):
            current_content.append(('hr', ''))
        elif line.startswith('|'):
            current_content.append(('table', line.strip()))
        elif line.strip():
            current_content.append(('text', line.strip()))

    if current_section:
        sections.append((current_section, current_content))

    return sections


def apply_bold_italic(paragraph, text):
    """Parse inline markdown bold/italic and add runs to paragraph."""
    # Split on **bold** patterns
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        else:
            # Handle *italic* within non-bold text
            sub_parts = re.split(r'(\*.*?\*)', part)
            for sub in sub_parts:
                if sub.startswith('*') and sub.endswith('*') and not sub.startswith('**'):
                    run = paragraph.add_run(sub[1:-1])
                    run.italic = True
                else:
                    paragraph.add_run(sub)


def md_to_docx(md_path, output_path):
    """Convert a markdown resume to ATS-friendly DOCX."""
    with open(md_path, 'r') as f:
        md_text = f.read()

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Set narrow margins for more content
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)

    sections = parse_markdown(md_text)

    for (stype, stitle), content in sections:
        if stype == 'h1':
            # Name / Title
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(stitle)
            run.bold = True
            run.font.size = Pt(18)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            p.space_after = Pt(2)

        elif stype == 'h2':
            # Section headers
            p = doc.add_paragraph()
            run = p.add_run(stitle.upper())
            run.bold = True
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)
            p.space_before = Pt(10)
            p.space_after = Pt(4)
            # Add a bottom border effect with a thin line
            p_border = doc.add_paragraph()
            p_border.space_before = Pt(0)
            p_border.space_after = Pt(4)
            run = p_border.add_run('_' * 90)
            run.font.size = Pt(2)
            run.font.color.rgb = RGBColor(0x1a, 0x1a, 0x2e)

        # Process content
        for ctype, ctext in content:
            if ctype == 'h3':
                p = doc.add_paragraph()
                apply_bold_italic(p, ctext)
                for run in p.runs:
                    run.font.size = Pt(11)
                    run.bold = True
                p.space_before = Pt(6)
                p.space_after = Pt(2)

            elif ctype == 'bullet':
                p = doc.add_paragraph(style='List Bullet')
                apply_bold_italic(p, ctext)
                for run in p.runs:
                    run.font.size = Pt(10)
                p.space_before = Pt(1)
                p.space_after = Pt(1)
                p.paragraph_format.left_indent = Inches(0.25)

            elif ctype == 'text':
                if ctext.startswith('**') and ctext.endswith('**'):
                    # Bold standalone line (like subtitle)
                    p = doc.add_paragraph()
                    apply_bold_italic(p, ctext)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.space_after = Pt(2)
                elif '|' in ctext and ctext.startswith('('):
                    # Contact info line
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    apply_bold_italic(p, ctext)
                    for run in p.runs:
                        run.font.size = Pt(10)
                    p.space_after = Pt(6)
                else:
                    p = doc.add_paragraph()
                    apply_bold_italic(p, ctext)
                    for run in p.runs:
                        run.font.size = Pt(10)
                    p.space_after = Pt(2)

            elif ctype == 'hr':
                pass  # Skip horizontal rules

            elif ctype == 'table':
                # Skip table formatting lines
                if ctext.startswith('|--') or ctext.startswith('| --'):
                    continue
                # Convert table rows to readable text
                cells = [c.strip() for c in ctext.split('|') if c.strip()]
                if cells:
                    p = doc.add_paragraph()
                    apply_bold_italic(p, ' | '.join(cells))
                    for run in p.runs:
                        run.font.size = Pt(9.5)
                    p.space_after = Pt(1)

    doc.save(output_path)
    return output_path


PDF_CSS = """
@page {
    size: letter;
    margin: 0.5in 0.6in;
}
body {
    font-family: 'Calibri', 'Helvetica Neue', Arial, sans-serif;
    font-size: 10.5pt;
    color: #333;
    line-height: 1.4;
}
h1 {
    text-align: center;
    font-size: 18pt;
    color: #1a1a2e;
    margin-bottom: 2pt;
}
h2 {
    font-size: 12pt;
    color: #1a1a2e;
    text-transform: uppercase;
    border-bottom: 1px solid #1a1a2e;
    padding-bottom: 3pt;
    margin-top: 10pt;
    margin-bottom: 4pt;
}
h3 {
    font-size: 11pt;
    margin-top: 6pt;
    margin-bottom: 2pt;
}
p {
    margin: 2pt 0;
    font-size: 10pt;
}
ul {
    margin: 2pt 0;
    padding-left: 0.25in;
}
li {
    font-size: 10pt;
    margin: 1pt 0;
}
hr {
    border: none;
    border-top: 1px solid #ccc;
    margin: 6pt 0;
}
/* Center the first paragraph (contact info) */
h1 + p {
    text-align: center;
    font-size: 10pt;
    margin-bottom: 6pt;
}
"""


def md_to_pdf(md_path, output_path):
    """Convert a markdown resume to PDF."""
    with open(md_path, 'r') as f:
        md_text = f.read()

    html_body = markdown2.markdown(md_text, extras=['tables', 'fenced-code-blocks'])
    full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>{PDF_CSS}</style></head>
<body>{html_body}</body></html>"""

    HTML(string=full_html).write_pdf(output_path)
    return output_path


COVER_LETTER_CSS = """
@page {
    size: letter;
    margin: 1in 1in;
}
body {
    font-family: 'Calibri', 'Helvetica Neue', Arial, sans-serif;
    font-size: 11pt;
    color: #222;
    line-height: 1.5;
}
.sender {
    margin-bottom: 24pt;
}
.sender .name {
    font-size: 16pt;
    font-weight: 700;
    color: #1a1a2e;
}
.sender .contact {
    font-size: 10pt;
    color: #555;
    margin-top: 2pt;
}
.sender .date {
    font-size: 10pt;
    color: #555;
    margin-top: 2pt;
}
.recipient {
    margin-bottom: 20pt;
    font-size: 11pt;
    line-height: 1.4;
}
.salutation {
    margin-bottom: 12pt;
    font-size: 11pt;
}
.body p {
    margin: 0 0 10pt 0;
    text-align: justify;
    font-size: 11pt;
}
.closing {
    margin-top: 20pt;
    margin-bottom: 4pt;
    font-size: 11pt;
}
.signature {
    font-weight: 700;
    font-size: 11pt;
    color: #1a1a2e;
}
"""


def md_to_cover_letter_pdf(md_path, output_path):
    """Convert a cover letter markdown to a professionally formatted PDF."""
    with open(md_path, 'r') as f:
        lines = f.read().strip().split('\n')

    # Group consecutive non-blank lines into blocks
    blocks = []
    current = []
    for line in lines:
        if line.strip() == '':
            if current:
                blocks.append(current)
                current = []
        else:
            current.append(line.strip())
    if current:
        blocks.append(current)

    html_parts = []

    # Block 0: Sender (name, contact info, date)
    if blocks:
        s = blocks[0]
        html_parts.append('<div class="sender">')
        if len(s) >= 1:
            html_parts.append(f'<div class="name">{s[0]}</div>')
        if len(s) >= 2:
            html_parts.append(f'<div class="contact">{s[1]}</div>')
        if len(s) >= 3:
            html_parts.append(f'<div class="date">{s[2]}</div>')
        html_parts.append('</div>')

    # Block 1: Recipient (company, position, optionally URL)
    if len(blocks) > 1:
        r = blocks[1]
        html_parts.append('<div class="recipient">')
        for line in r:
            if not line.startswith('http'):
                html_parts.append(f'<div>{line}</div>')
        html_parts.append('</div>')

    # Remaining blocks: body, closing, signature
    body_blocks = blocks[2:]
    for i, block in enumerate(body_blocks):
        text = ' '.join(block)
        is_last = (i == len(body_blocks) - 1)
        is_second_last = (i == len(body_blocks) - 2)

        if is_last and len(block) == 1:
            # Final single-line block = signature
            html_parts.append(f'<div class="signature">{text}</div>')
        elif is_second_last and text.rstrip(',').lower() in (
                'sincerely', 'best regards', 'regards', 'best', 'respectfully'):
            html_parts.append(f'<div class="closing">{text}</div>')
        elif text.startswith('Dear '):
            html_parts.append(f'<div class="salutation">{text}</div>')
            html_parts.append('<div class="body">')
        else:
            html_parts.append(f'<p>{text}</p>')

    # Close body div if we opened one
    if any('class="body"' in p for p in html_parts):
        # Insert closing </div> before the closing/signature
        for i in range(len(html_parts) - 1, -1, -1):
            if 'class="closing"' in html_parts[i]:
                html_parts.insert(i, '</div>')
                break

    html_body = '\n'.join(html_parts)
    full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8"><style>{COVER_LETTER_CSS}</style></head>
<body>{html_body}</body></html>"""

    HTML(string=full_html).write_pdf(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description='Convert markdown resumes to ATS-friendly DOCX')
    parser.add_argument('input', nargs='?', help='Input markdown file path')
    parser.add_argument('--all', action='store_true', help='Convert all resumes in Resumes/ directory')
    parser.add_argument('--output', default=None, help='Output directory (default: same as input)')
    parser.add_argument('--pdf', action='store_true', help='Convert to PDF instead of DOCX')
    args = parser.parse_args()

    base_dir = Path(__file__).parent
    resume_dir = base_dir / 'Resumes'
    output_dir = Path(args.output) if args.output else resume_dir

    if not output_dir.exists():
        output_dir.mkdir(parents=True)

    fmt = 'pdf' if args.pdf else 'docx'
    ext = f'.{fmt}'
    converter = md_to_pdf if args.pdf else md_to_docx

    if args.all:
        md_files = sorted(resume_dir.glob('*.md'))
        if not md_files:
            print("No markdown files found in Resumes/ directory")
            sys.exit(1)

        print(f"Converting {len(md_files)} resumes to {fmt.upper()}...\n")
        for md_file in md_files:
            out_name = md_file.stem + ext
            output_path = output_dir / out_name
            try:
                converter(str(md_file), str(output_path))
                print(f"  ✓ {md_file.name} → {out_name}")
            except Exception as e:
                print(f"  ✗ {md_file.name} FAILED: {e}")

        print(f"\nDone! {len(md_files)} resumes converted to {output_dir}/")

    elif args.input:
        md_file = Path(args.input)
        if not md_file.exists():
            print(f"File not found: {md_file}")
            sys.exit(1)

        out_name = md_file.stem + ext
        output_path = output_dir / out_name
        converter(str(md_file), str(output_path))
        print(f"Converted: {md_file.name} → {output_path}")

    else:
        parser.print_help()
        sys.exit(1)


if __name__ == '__main__':
    main()
